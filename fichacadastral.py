import os
import PySimpleGUI as sg
from docx import Document
from utils import load_excel_data

def validate_paths(input_path, output_path, docx_file):
    if not os.path.exists(input_path):
        sg.popup_error("Caminho de entrada inválido.")
        return False
    if not os.path.exists(docx_file):
        sg.popup_error("Arquivo DOCX não encontrado.")
        return False
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    return True

def replace_text_in_document(doc, info):
    text_replaced = False
    for paragraph in doc.paragraphs:
        for placeholder, value in info.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                text_replaced = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in info.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        text_replaced = True
    return text_replaced

def replace_text_in_fichacadastral(combo_value, docx_file, output_path, input_path):
    if not validate_paths(input_path, output_path, docx_file):
        return

    excel_file_path = os.path.join(input_path, "NLcadastro.xlsm")
    ws, names = load_excel_data(excel_file_path)
    if not ws:
        sg.popup_error("Falha ao carregar dados do Excel.")
        return

    info = None
    for row in ws.iter_rows(min_row=2, max_col=ws.max_column, values_only=True):
        if row[5] == combo_value:
            info = {
                '#email': row[1] or '',
                '#nomeempresa': row[2] or '',
                '#cnpj': row[3] or '',
                '#sede': row[4] or '',
                '#nome': row[5] or '',
                '#tel': row[6] or '',
                '#estadocivil': row[7] or '',
                '#profissao': row[8] or '',
                '#cpf': row[9] or '',
                '#rg': row[11] or '',
                '#mae': row[12] or '',
                '#pai': row[13] or '',
                '#nascimento': row[14] or '',
                '#endereco': row[15] or '',
                '#num': row[16] or '',
                '#compl': row[17] or '',
                '#cep': row[18] or '',
                '#cidade': row[19] or '',
                '#estado': row[20] or '',
                '#localnasc': row[21] or '',
                '#nacionalidade': row[22] or '',
                '#datapag': row[23] or '',
                '#formpag': row[24] or '',
                '#meiopag': row[25] or '',
            }
            break

    if not info:
        sg.popup("Nome não encontrado.")
        return

    try:
        doc = Document(docx_file)
    except Exception as e:
        sg.popup_error(f"Erro ao abrir o arquivo DOCX: {e}")
        return

    text_replaced = replace_text_in_document(doc, info)

    if text_replaced:
        output_file = os.path.join(output_path, f"Ficha_Cadastral_{combo_value.replace(' ', '_')}.docx")
        try:
            doc.save(output_file)
            sg.popup("Documento gerado com sucesso!")
        except Exception as e:
            sg.popup_error(f"Erro ao salvar o documento: {e}")
    else:
        sg.popup_error("Nenhuma substituição de texto foi realizada.")
